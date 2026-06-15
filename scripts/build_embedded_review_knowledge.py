from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("output/doc")
MD_PATH = OUT_DIR / "2026春嵌入式系统原理复习纲要-完全版知识体系.md"
DOCX_PATH = OUT_DIR / "2026春嵌入式系统原理复习纲要-完全版知识体系.docx"


chapters = [
    {
        "title": "一、Linux 操作系统基础",
        "lead": "嵌入式开发通常不是直接在目标开发板上写代码，而是在 PC 上编写、交叉编译、下载到开发板运行。Linux 是最常见的开发环境，也是很多机器人与嵌入式设备的运行环境。",
        "sections": [
            (
                "嵌入式系统是什么",
                [
                    "嵌入式系统是嵌入到具体设备内部、完成特定功能的计算机系统。它不像普通电脑那样追求通用性，而是服务于某个明确场景，例如智能手表、路由器、工业控制器、机器人、汽车控制单元。",
                    "一个完整嵌入式系统通常包括硬件、启动程序、操作系统或裸机程序、驱动、应用程序和外设接口。硬件负责采集和执行，软件负责控制逻辑和任务调度。",
                    "学习时要先建立一个总图：传感器采集数据，处理器运行控制程序，执行器完成动作，操作系统负责管理 CPU、内存、文件、进程、网络和设备。",
                ],
            ),
            (
                "嵌入式系统的主要特点",
                [
                    "专用性强：目标明确，常围绕一个任务设计，例如机器人导航、温度控制、门禁识别。专用性意味着功能可以更集中，硬件和软件都可以围绕目标优化。",
                    "资源受限：CPU 主频、内存、存储、电池、网络带宽往往有限。程序要注意效率，避免无意义的内存占用和复杂计算。",
                    "实时性要求高：系统必须在规定时间内做出响应。例如刹车控制不能晚，机器人避障不能等。实时性强调“按时完成”，不只是“算得快”。",
                    "功耗敏感：许多设备靠电池运行，软件要考虑休眠、低功耗模式、降低外设频率等策略。",
                    "高可靠性与稳定性：嵌入式设备常长时间无人值守运行，死机、内存泄漏、异常重启都会造成风险。",
                    "软硬件协同设计：软件必须理解硬件接口、寄存器、传感器协议和执行器特性，硬件也要为软件控制提供支持。",
                    "系统内核精简：嵌入式操作系统一般只保留必要组件，减少体积和启动时间。",
                    "开发环境特殊：常见流程是宿主机开发、交叉编译、烧录或部署到目标板调试。",
                    "需要 BSP 支持：BSP 是 Board Support Package，板级支持包，负责让操作系统适配具体开发板硬件，包括启动、时钟、引脚、驱动等。",
                    "成本敏感：实际产品会严格控制芯片、存储、传感器和制造成本，软件设计也要服务于成本目标。",
                ],
            ),
            (
                "嵌入式操作系统的主要特点",
                [
                    "实时性：嵌入式 OS 常需要提供可预测的任务调度、定时器、中断响应和优先级机制。",
                    "可裁剪：可以按需要保留文件系统、网络协议栈、图形界面、驱动等模块，降低系统体积。",
                    "资源受限：操作系统本身要轻量，不能消耗过多内存和 CPU。",
                    "高可靠：强调长期运行、异常恢复、看门狗、日志和故障隔离。",
                    "低功耗：支持设备休眠、动态频率调整、外设电源管理。",
                    "硬件相关：嵌入式 OS 与具体处理器架构、开发板外设、驱动适配高度相关。",
                    "交叉开发：开发机和运行设备通常不是同一种架构，例如在 x86 电脑上编译 ARM 开发板程序。",
                    "固化运行：程序常被烧录到 Flash、eMMC、SD 卡或镜像中，开机自动启动。",
                ],
            ),
            (
                "Ubuntu 的两种工作模式",
                [
                    "图形化界面模式：通过窗口、鼠标和菜单操作，适合初学者浏览文件、编辑文本和观察系统状态，但占用资源较多。",
                    "命令行工作模式：通过终端输入命令操作。它需要记忆命令，但资源占用低、效率高、便于远程连接和自动化脚本。",
                    "在 ROS 和嵌入式开发中，终端非常重要。编译、启动节点、查看话题、安装软件、配置环境变量都主要依赖命令行。",
                ],
            ),
            (
                "Ubuntu 在线安装常用命令",
                [
                    "`sudo apt update`：同步软件源列表，意思是让本机知道软件仓库里有哪些软件和版本。它不会真正升级已安装软件。",
                    "`sudo apt upgrade`：根据已经同步的软件列表，升级本机已安装的软件包。",
                    "`sudo apt install 包名`：安装指定软件包。例如安装 ROS 工具、编译器、库文件。",
                    "`sudo apt remove 包名`：卸载软件包。",
                    "`apt search 关键字`：搜索软件包。",
                    "update 和 upgrade 的核心区别：update 更新“清单”，upgrade 更新“软件本体”。考试中常把二者混淆。",
                ],
            ),
        ],
        "review": [
            "能说清嵌入式系统和普通 PC 的区别。",
            "能解释实时性、资源受限、BSP、交叉开发。",
            "能区分 apt update 与 apt upgrade。",
        ],
    },
    {
        "title": "二、ROS 操作系统简介",
        "lead": "ROS 常被称为机器人操作系统，但它本质上不是传统意义上的操作系统内核，而是一套机器人软件框架和工具生态。它运行在 Linux 等系统之上，帮助机器人程序拆成多个节点并互相通信。",
        "sections": [
            (
                "ROS 全称、定义及起源",
                [
                    "ROS 全称是 Robot Operating System，即机器人操作系统。",
                    "ROS 是一个分布式机器人软件框架，提供通信机制、包管理、编译系统、可视化工具、仿真工具和大量机器人算法功能包。",
                    "ROS 起源于机器人研究和开源社区，目标是复用机器人软件，避免每个团队都从零编写驱动、通信、导航、视觉和仿真模块。",
                    "理解 ROS 时可以把它看成“机器人软件积木系统”：每个功能包解决一类问题，每个节点完成一个功能，节点之间通过消息协作。",
                ],
            ),
            (
                "ROS 的五个特点",
                [
                    "点对点分布式：多个节点可以运行在同一台机器或不同机器上，通过网络通信。",
                    "多语言支持：常用 C++ 和 Python，也支持其他语言绑定。",
                    "工具丰富：提供 roscore、rosrun、roslaunch、rviz、rqt_graph、rosbag 等工具。",
                    "功能包生态庞大：导航、建图、视觉、机械臂、传感器驱动等都有大量开源包。",
                    "开源与复用：ROS 鼓励把代码组织成功能包，便于分享、移植和二次开发。",
                ],
            ),
            (
                "ROS 架构",
                [
                    "文件系统层：包括工作空间、src 目录、功能包 package、CMakeLists.txt、package.xml、msg、srv、launch 等。",
                    "计算图层：运行时由节点、主题、服务、动作、参数服务器、节点管理器组成。",
                    "社区层：包括 ROS Wiki、发行版、开源功能包和维护工具。",
                    "初学者最重要的是计算图层：节点是程序，主题/服务/动作/参数是节点之间协作的方式。",
                ],
            ),
            (
                "ROS Noetic 与 Ubuntu 版本",
                [
                    "ROS Noetic Ninjemys 是 ROS 1 的重要发行版，主要基于 Ubuntu 20.04 LTS。",
                    "考试中如果问 Noetic 对应的 Ubuntu 版本，答案是 Ubuntu 20.04 LTS。",
                    "ROS 版本通常与 Ubuntu 版本绑定，因为底层依赖、编译器、Python 版本和系统库不同。",
                ],
            ),
            (
                "ROS 编译系统与常用命令",
                [
                    "ROS 1 常用 catkin 编译系统。catkin 基于 CMake，用来管理功能包依赖、生成消息/服务代码、编译可执行文件。",
                    "`catkin_make`：在工作空间根目录编译整个工作空间。",
                    "`source devel/setup.bash`：让当前终端识别工作空间里的功能包、节点和消息类型。每次打开新终端通常都要 source。",
                    "`rospack find 包名`：查找功能包路径。",
                    "`roscd 包名`：进入功能包目录。",
                    "`rosrun 包名 节点名`：运行某个包里的一个节点。",
                    "`roslaunch 包名 文件.launch`：按 launch 文件一次启动多个节点和参数。",
                ],
            ),
            (
                "ROS 工作空间与文件夹构成",
                [
                    "典型 catkin 工作空间包含 `src`、`build`、`devel` 三个主要目录。",
                    "`src`：放源代码和功能包，是开发者主要编辑的地方。",
                    "`build`：编译中间文件目录，一般不手动修改。",
                    "`devel`：开发环境输出目录，包含生成的可执行文件、库、消息代码和 setup 脚本。",
                    "功能包内部常见文件：`package.xml` 描述包名、版本、依赖；`CMakeLists.txt` 描述编译规则；`src` 放 C++ 源文件；`scripts` 放 Python 脚本；`msg` 放自定义消息；`srv` 放自定义服务；`launch` 放启动配置。",
                ],
            ),
            (
                "节点、节点管理器与启动",
                [
                    "节点 node 是 ROS 中执行具体功能的进程。例如摄像头采集节点、速度控制节点、路径规划节点。",
                    "节点管理器 ROS Master 由 `roscore` 启动，负责节点注册、名称解析和连接信息交换。它不负责转发所有数据，只帮助节点找到彼此。",
                    "启动顺序通常是先运行 `roscore`，再用 `rosrun` 或 `roslaunch` 启动节点。",
                    "如果使用 `roslaunch`，它会在需要时自动启动 roscore，因此不一定要单独手动启动。",
                ],
            ),
        ],
        "review": [
            "能解释 ROS 不是传统 OS 内核，而是机器人软件框架。",
            "能画出工作空间 src/build/devel 结构。",
            "能说明 roscore、rosrun、roslaunch、catkin_make 的作用。",
        ],
    },
    {
        "title": "三、ROS 通信方式",
        "lead": "ROS 的核心是通信。机器人系统由很多节点组成，节点之间必须交换传感器数据、控制命令、配置参数和任务结果。",
        "sections": [
            (
                "异步编程模式和思想",
                [
                    "同步调用像打电话：发起方必须等对方回复，才能继续下一步。",
                    "异步通信像发消息：发送方发出信息后可以继续做别的事，接收方在合适时机处理。",
                    "机器人系统常需要异步，因为传感器数据持续到来，控制程序不能因等待某个结果而卡死。",
                    "ROS 主题通信就是典型异步模式：发布者不断发布消息，订阅者通过回调函数处理收到的数据。",
                ],
            ),
            (
                "ROS 四种通信方式总览",
                [
                    "主题 Topic：多对多、异步、连续数据流，适合传感器数据、速度命令、状态发布。例如 `/cmd_vel`、`/scan`、`/camera/image_raw`。",
                    "服务 Service：一问一答、同步、短时请求，适合需要明确返回值的操作。例如查询当前状态、触发一次保存地图。",
                    "动作 Action：异步、可反馈、可取消，适合耗时任务。例如导航到目标点、机械臂抓取任务。",
                    "参数服务器 Parameter Server：保存全局配置参数，适合机器人名称、控制频率、阈值、文件路径等低频配置。",
                    "选择通信方式的关键：连续数据用主题；短请求用服务；耗时任务用动作；全局配置用参数。",
                ],
            ),
            (
                "主题编程流程",
                [
                    "发布者流程：初始化节点，创建 NodeHandle，创建 Publisher，构造消息，按频率 publish，调用 spinOnce 或循环等待。",
                    "订阅者流程：初始化节点，创建 NodeHandle，创建 Subscriber，指定回调函数，在回调函数中处理消息，调用 spin 进入循环。",
                    "核心函数包括 `ros::init`、`ros::NodeHandle`、`advertise`、`subscribe`、`publish`、`ros::Rate`、`spin`、`spinOnce`。",
                    "小乌龟速度控制的本质是向 `/turtle1/cmd_vel` 主题发布 `geometry_msgs::Twist` 消息，turtlesim 节点订阅后改变乌龟运动。",
                ],
            ),
            (
                "geometry_msgs::Twist 类型",
                [
                    "`Twist` 用来表示线速度和角速度，包含两个三维向量：`linear` 和 `angular`。",
                    "`linear.x` 常表示机器人前后方向速度，单位通常是 m/s；`angular.z` 常表示绕垂直轴旋转角速度，单位通常是 rad/s。",
                    "差速小车或小乌龟常用 `linear.x` 控制前进，`angular.z` 控制转弯。其他分量在平面移动中通常设为 0。",
                    "典型代码：`vel_msg.linear.x = 0.5; vel_msg.angular.z = 1.0; pub.publish(vel_msg);`。",
                ],
            ),
            (
                "rqt_graph 的作用",
                [
                    "`rqt_graph` 是 ROS 计算图可视化工具，可以显示节点、主题以及发布订阅关系。",
                    "它适合检查节点是否启动、话题是否连接正确、命名空间是否写错。",
                    "如果程序没有反应，先用 `rostopic list` 和 `rqt_graph` 看发布者与订阅者是否连上，是很常见的排错方法。",
                ],
            ),
            (
                "CMakeLists.txt 编译规则",
                [
                    "`add_executable(节点名 src/文件.cpp)`：把源文件编译成可运行节点。",
                    "`add_dependencies(节点名 ${${PROJECT_NAME}_EXPORTED_TARGETS} ${catkin_EXPORTED_TARGETS})`：确保消息、服务等自动生成代码先生成，再编译节点。",
                    "`target_link_libraries(节点名 ${catkin_LIBRARIES})`：把 ROS/catkin 相关库链接到可执行文件。",
                    "如果使用自定义 msg 或 srv，还需要 `add_message_files`、`add_service_files`、`generate_messages` 和 `catkin_package` 中的依赖配置。",
                ],
            ),
            (
                "spin 与 spinOnce",
                [
                    "`ros::spin()` 会进入循环，持续等待并处理回调，常用于订阅者节点。调用后一般不会往下执行。",
                    "`ros::spinOnce()` 只处理一次当前回调队列，然后立即返回，常与 `while(ros::ok())` 和 `ros::Rate` 配合，用于既要发布又要处理回调的节点。",
                    "简单记法：只等回调用 spin；循环里还要干别的事用 spinOnce。",
                ],
            ),
            (
                "创建和使用自定义服务",
                [
                    "在功能包中新建 `srv` 目录，创建 `.srv` 文件。文件中用 `---` 分隔请求和响应。例如请求两个整数，响应它们的和。",
                    "修改 `package.xml`，加入 `message_generation` 和 `message_runtime` 等依赖。",
                    "修改 `CMakeLists.txt`，配置 `add_service_files` 和 `generate_messages`。",
                    "编译后 source 工作空间，服务类型才能被识别。",
                    "服务端用 `advertiseService` 注册服务，回调函数接收 Request、填写 Response。",
                    "客户端用 `serviceClient` 创建客户端，构造请求后调用 `client.call(srv)`。",
                ],
            ),
            (
                "动作库通信概述",
                [
                    "动作 Action 适合执行时间较长、需要过程反馈、可能取消的任务。",
                    "动作包含 Goal、Feedback、Result 三类信息。Goal 是目标，Feedback 是执行过程状态，Result 是最终结果。",
                    "导航中发送目标点给 move_base 就是典型动作通信：机器人移动过程中持续反馈状态，成功或失败后返回结果。",
                ],
            ),
            (
                "参数服务器维护方式",
                [
                    "参数服务器保存键值形式的配置。常见参数类型包括整数、浮点数、字符串、布尔值、列表和字典。",
                    "`rosparam list` 查看参数列表；`rosparam get 参数名` 获取参数；`rosparam set 参数名 值` 设置参数；`rosparam delete 参数名` 删除参数。",
                    "`rosparam load 文件.yaml` 从 YAML 文件加载参数；`rosparam dump 文件.yaml` 导出参数。",
                    "代码中可以通过 `nh.getParam`、`nh.setParam`、`param` 等函数访问参数。",
                ],
            ),
        ],
        "review": [
            "能比较 Topic、Service、Action、Parameter Server。",
            "能写出发布者/订阅者基本流程。",
            "能解释 Twist、rqt_graph、spin/spinOnce、CMakeLists 核心函数。",
        ],
    },
    {
        "title": "四、ROS 实用工具",
        "lead": "ROS 工具帮助我们启动系统、观察数据、调试坐标关系、查看仿真和可视化结果。掌握工具比死记代码更能提高排错能力。",
        "sections": [
            (
                "TF 坐标变换概述",
                [
                    "机器人上有很多坐标系，例如 map、odom、base_link、laser、camera_link。不同传感器数据必须转换到统一坐标系才能融合。",
                    "TF 是 ROS 中管理坐标变换的工具库，用来维护坐标系之间的位置和姿态关系。",
                    "例如激光雷达测到障碍物在 laser 坐标系下的位置，导航算法可能需要把它转换到 map 或 base_link 坐标系下。",
                    "TF 变换包括平移和旋转。旋转常用四元数表示，因为四元数比欧拉角更适合连续三维旋转计算。",
                ],
            ),
            (
                "使用 TF 功能包的步骤",
                [
                    "确定机器人需要哪些坐标系，例如 map、odom、base_link、camera、laser。",
                    "发布坐标变换。静态关系可用 `static_transform_publisher`，动态关系由里程计、定位或驱动节点持续发布。",
                    "在程序中创建 TF 监听器，查询两个坐标系之间的变换。",
                    "把点、姿态或传感器数据转换到目标坐标系，再交给后续算法处理。",
                ],
            ),
            (
                "TF 树的作用及查看方式",
                [
                    "TF 树把所有坐标系组织成树状结构，每条边表示一个坐标变换。",
                    "TF 树要求坐标系之间关系清晰，通常不能出现多个父节点导致歧义。",
                    "查看方式：`rosrun rqt_tf_tree rqt_tf_tree` 可以生成坐标树图；`rosrun tf view_frames` 可以导出坐标关系；`rosrun tf tf_echo frame1 frame2` 可以查看实时变换。",
                ],
            ),
            (
                "launch 文件编写规则与启动",
                [
                    "launch 文件是 XML 格式，根标签是 `<launch>`。",
                    "`<node>` 用来启动节点，常用属性包括 `pkg`、`type`、`name`、`output`。",
                    "`<param>` 设置单个参数，`<rosparam>` 加载 YAML 参数文件。",
                    "`<include>` 引入其他 launch 文件，便于复用复杂启动配置。",
                    "`<arg>` 定义可传入参数，适合在同一个 launch 中切换配置。",
                    "启动命令：`roslaunch 包名 文件.launch`。",
                ],
            ),
            (
                "rosrun 与 roslaunch 的区别",
                [
                    "`rosrun` 一次启动一个节点，适合简单测试。",
                    "`roslaunch` 可以一次启动多个节点、设置参数、加载配置、指定命名空间，还能自动启动 roscore。",
                    "机器人系统通常由多个节点组成，因此正式运行更常用 roslaunch。",
                ],
            ),
            (
                "roscore 与 roslaunch 的关系",
                [
                    "`roscore` 启动 ROS Master、参数服务器和日志相关组件。",
                    "`roslaunch` 启动前会检查 roscore 是否存在；如果不存在，通常会自动启动一个。",
                    "所以使用 roslaunch 时不一定需要提前手动运行 roscore，但理解 roscore 的作用仍然很重要。",
                ],
            ),
            (
                "Gazebo 软件",
                [
                    "Gazebo 是机器人三维物理仿真平台，可以模拟机器人、环境、重力、碰撞、传感器和执行器。",
                    "功能：加载机器人模型，模拟激光雷达、相机、IMU，进行运动控制、导航、抓取和多机器人实验。",
                    "特点：有物理引擎，关注真实运动和环境交互。",
                    "应用场景：没有真实机器人时测试算法，降低实验成本和风险；在部署前验证控制和导航逻辑。",
                ],
            ),
            (
                "下载开源项目到 ROS 工作空间",
                [
                    "进入工作空间的 `src` 目录。",
                    "使用 `git clone 项目地址` 下载源码。",
                    "回到工作空间根目录执行 `catkin_make` 编译。",
                    "执行 `source devel/setup.bash` 刷新环境。",
                    "根据项目说明安装依赖，常用 `rosdep install --from-paths src --ignore-src -r -y`。",
                ],
            ),
            (
                "Rviz 软件",
                [
                    "Rviz 是 ROS 的三维可视化工具，主要显示机器人状态和传感器数据，而不是做真实物理仿真。",
                    "可以显示 TF 坐标系、机器人模型、激光点云、地图、路径、导航目标、相机图像等。",
                    "Rviz 适合观察算法输入输出是否正确，例如地图是否显示、激光是否对齐、路径是否合理。",
                ],
            ),
            (
                "Gazebo 与 Rviz 的主要区别",
                [
                    "Gazebo 负责仿真世界和物理交互，回答“机器人在虚拟环境中会怎样运动”。",
                    "Rviz 负责可视化 ROS 数据，回答“当前 ROS 系统发布的数据是什么样”。",
                    "Gazebo 像实验场地，Rviz 像调试仪表盘。两者常同时使用。",
                ],
            ),
        ],
        "review": [
            "能解释 TF 坐标树、launch 文件、rosrun/roslaunch。",
            "能区分 Gazebo 和 Rviz。",
            "能描述下载并编译 ROS 开源包的基本步骤。",
        ],
    },
    {
        "title": "五、机器人建模与运动仿真",
        "lead": "机器人仿真首先要让计算机知道机器人长什么样、有哪些部件、关节如何运动、传感器安装在哪里。这通常由 URDF 完成。",
        "sections": [
            (
                "URDF 的结构和描述",
                [
                    "URDF 全称 Unified Robot Description Format，是 ROS 中描述机器人模型的 XML 格式。",
                    "URDF 描述机器人连杆、关节、外观、碰撞几何、惯性参数和坐标关系。",
                    "典型结构包括 `<robot>` 根标签，内部包含多个 `<link>` 和 `<joint>`。",
                    "URDF 可以被 robot_state_publisher、Rviz、Gazebo 等工具使用，用于显示模型和建立 TF 关系。",
                ],
            ),
            (
                "连杆 link 的含义",
                [
                    "连杆表示机器人上的刚体部件，例如底盘、轮子、机械臂某一节、摄像头支架。",
                    "link 可以包含 visual、collision、inertial 三类信息。",
                    "visual 描述外观，用于显示；collision 描述碰撞几何，用于物理检测；inertial 描述质量和惯性，用于动力学仿真。",
                    "初学时可以先理解为：link 是机器人身体的一块固定部件。",
                ],
            ),
            (
                "关节 joint 的含义",
                [
                    "关节描述两个 link 之间的连接关系和运动方式。",
                    "常见关节类型：fixed 固定关节；revolute 有限角度旋转关节；continuous 连续旋转关节；prismatic 平移关节；floating 浮动关节；planar 平面关节。",
                    "joint 通常包含 parent、child、origin、axis、limit 等信息。parent 和 child 指明连接的两个连杆；axis 指明运动轴；limit 指明运动范围。",
                    "机器人模型本质上是 link 和 joint 组成的树。",
                ],
            ),
            (
                "ROS 支持的传感器类型",
                [
                    "激光雷达 LiDAR：输出距离扫描或点云，适合建图、定位、避障。二维雷达常输出 LaserScan，三维雷达常输出 PointCloud2。",
                    "相机 Camera：输出图像，适合目标检测、跟随、识别、视觉 SLAM。普通相机输出 RGB 图像，深度相机还能输出深度图或点云。",
                    "IMU：惯性测量单元，测量角速度、加速度和姿态相关数据，适合姿态估计和运动状态估计。",
                    "编码器 Encoder：测量轮子或关节转动，用于里程计估计和闭环控制。",
                    "GPS：提供全球位置，适合室外机器人定位，但室内通常不可用或精度差。",
                    "超声波/红外传感器：成本低，适合近距离测距和简单避障，但精度和抗干扰能力有限。",
                    "触碰/力传感器：检测接触、压力或抓取力，适合机械臂和服务机器人交互。",
                ],
            ),
        ],
        "review": [
            "能说清 URDF、link、joint 的关系。",
            "能列举常见传感器及其适用场景。",
            "能理解仿真模型为什么需要 visual、collision、inertial。",
        ],
    },
    {
        "title": "六、机器人建图与导航仿真应用",
        "lead": "机器人导航要解决三个问题：我在哪里、地图是什么样、我怎样从当前位置到目标点。SLAM、定位、路径规划和运动控制共同完成这件事。",
        "sections": [
            (
                "SLAM 建图仿真定义和主要功能",
                [
                    "SLAM 全称 Simultaneous Localization and Mapping，即同步定位与建图。",
                    "机器人在未知环境中一边估计自身位置，一边构建环境地图。",
                    "建图仿真的意义是在 Gazebo 等仿真环境中验证算法，不需要真实机器人和真实场地。",
                    "SLAM 主要功能包括传感器数据采集、位姿估计、地图更新、回环检测和地图保存。",
                    "二维移动机器人常用激光雷达 SLAM，生成二维栅格地图。栅格地图中每个格子表示占用、空闲或未知。",
                ],
            ),
            (
                "Navigation 经典架构",
                [
                    "ROS Navigation 是 ROS 1 中经典二维导航框架，主要面向差速或全向移动机器人。",
                    "典型输入包括地图、机器人传感器数据、里程计、TF、导航目标点。",
                    "典型输出是速度命令 `/cmd_vel`，驱动底盘运动。",
                    "导航系统通常包括地图服务器、定位、全局代价地图、局部代价地图、全局规划器、局部规划器和恢复行为。",
                ],
            ),
            (
                "Navigation 设计理念和常用功能包",
                [
                    "`gmapping`：基于激光雷达和里程计进行 SLAM 建图，生成二维栅格地图。",
                    "`map_server`：加载和保存地图。常见文件包括 `.pgm` 地图图片和 `.yaml` 地图配置。",
                    "`amcl`：自适应蒙特卡洛定位，在已有地图中根据激光和里程计估计机器人位姿。",
                    "`move_base`：导航核心协调器，接收目标点，调用全局规划器和局部规划器，输出速度命令。",
                    "设计理念是模块化：建图、定位、规划、控制各司其职，通过 ROS 消息和 TF 连接。",
                ],
            ),
            (
                "全局规划器的作用",
                [
                    "全局规划器根据全局地图和目标点规划从起点到终点的大路径。",
                    "它关注整体路线，例如绕过墙、选择走廊、从房间到大厅。",
                    "常见算法思想包括 Dijkstra、A* 等图搜索算法。",
                    "全局路径不一定直接可执行，因为现场可能有临时障碍物，所以还需要局部规划器。",
                ],
            ),
            (
                "局部规划器的作用",
                [
                    "局部规划器根据当前机器人状态、局部代价地图和全局路径，实时生成可执行速度命令。",
                    "它负责处理动态障碍物，例如突然出现的行人、移动物体、局部地图变化。",
                    "局部规划器需要考虑机器人运动学约束、速度限制、加速度限制、避障距离和目标方向。",
                    "可以把全局规划器理解为“规划路线”，局部规划器理解为“边走边调整方向和速度”。",
                ],
            ),
            (
                "move_base 输入和默认输出消息格式",
                [
                    "move_base 常用动作接口接收 `move_base_msgs/MoveBaseAction`，目标通常是 `geometry_msgs/PoseStamped`，包含目标位置和姿态，以及所在坐标系。",
                    "move_base 依赖输入：地图、TF、里程计 `/odom`、激光 `/scan` 或点云、导航目标、代价地图参数。",
                    "默认输出通常是速度命令 `/cmd_vel`，消息类型为 `geometry_msgs/Twist`。",
                    "当 move_base 无法前进时，可能触发恢复行为，例如清除代价地图、原地旋转等。",
                ],
            ),
            (
                "经典局部规划器举例",
                [
                    "DWA Local Planner：Dynamic Window Approach，动态窗口法，在速度空间中采样可行速度，选择兼顾目标、路径和避障的速度。",
                    "Trajectory Rollout：模拟多条短期轨迹，选择评分最优的一条。",
                    "TEB Local Planner：Timed Elastic Band，把路径看成带时间约束的弹性带，适合处理动态障碍和复杂约束。",
                    "考试中不一定要求推导算法，但要知道它们解决的是局部避障和实时速度生成问题。",
                ],
            ),
        ],
        "review": [
            "能解释 SLAM、Navigation、gmapping、map_server、amcl、move_base。",
            "能区分全局规划器和局部规划器。",
            "能说出 move_base 输入与 `/cmd_vel` 输出。",
        ],
    },
    {
        "title": "七、机器人平面视觉检测仿真应用",
        "lead": "机器人视觉把摄像头图像变成可计算的信息，再把识别结果转化为控制命令。平面视觉常用 OpenCV 处理二维图像。",
        "sections": [
            (
                "获取机器人图像并转换为 OpenCV 格式",
                [
                    "ROS 中相机图像常用 `sensor_msgs/Image` 消息发布。",
                    "OpenCV 使用 `cv::Mat` 表示图像矩阵。ROS 图像和 OpenCV 图像格式不同，需要桥接。",
                    "`cv_bridge` 是 ROS 与 OpenCV 之间的转换工具，可以把 `sensor_msgs/Image` 转成 `cv::Mat`，也可以反向转换。",
                    "典型流程：订阅相机话题；在回调函数中用 cv_bridge 转换；用 OpenCV 处理；显示、发布结果或生成控制命令。",
                ],
            ),
            (
                "OpenCV 主要函数",
                [
                    "`cvCopy`：旧版 C 接口中的图像复制函数，用于把一幅图像内容复制到另一幅图像。现代 C++ 中更常用 `copyTo` 或 `clone`。",
                    "`cvtColor`：颜色空间转换函数。例如把 BGR 图像转成灰度图或 HSV 图。常见写法是 `cv::cvtColor(src, dst, cv::COLOR_BGR2HSV)`。",
                    "`inRange`：根据阈值提取指定颜色范围，常用于 HSV 颜色目标检测。",
                    "`erode` 和 `dilate`：腐蚀和膨胀，用于去噪、填补空洞、改善二值图。",
                    "`findContours`：寻找二值图中的轮廓，用于定位目标形状和区域。",
                    "`moments`：计算图像矩，可以得到目标区域中心点。",
                ],
            ),
            (
                "目标跟随闭环控制流程",
                [
                    "第一步：相机节点发布图像。",
                    "第二步：视觉节点订阅图像并转换为 OpenCV 格式。",
                    "第三步：根据颜色、形状或分类器检测目标，得到目标在图像中的位置和大小。",
                    "第四步：计算误差。例如目标中心与图像中心的横向偏差。",
                    "第五步：根据误差生成控制量。例如偏左就向左转，目标太小就前进。",
                    "第六步：发布 `geometry_msgs::Twist` 到 `/cmd_vel`，机器人执行运动。",
                    "闭环控制的关键是反馈：每一帧图像都会重新计算误差并调整速度。",
                ],
            ),
            (
                "色彩空间和 HSV 的意义",
                [
                    "常见 RGB/BGR 色彩空间用红、绿、蓝三个通道表示颜色，适合显示，但不一定适合稳定识别颜色。",
                    "HSV 用 Hue 色相、Saturation 饱和度、Value 明度表示颜色。",
                    "H 表示颜色种类，例如红、黄、绿；S 表示颜色纯度；V 表示亮暗。",
                    "HSV 的好处是把颜色信息和亮度信息相对分离，做颜色阈值检测时通常比 RGB 更稳定。",
                    "例如检测红色物体时，可以主要限制 H 的范围，再用 S 和 V 过滤灰暗或过亮区域。",
                ],
            ),
            (
                "Haar 级联分类器",
                [
                    "Haar 级联分类器是一种传统目标检测方法，常用于人脸检测等任务。",
                    "它使用 Haar-like 特征描述图像局部明暗差异，通过 AdaBoost 选择有效特征，并用级联结构快速排除非目标区域。",
                    "级联的意义是前几层快速过滤大量无关窗口，后面层再做更精细判断，从而提高检测速度。",
                    "它的优点是速度快、实现成熟；缺点是对姿态、光照、遮挡和复杂背景鲁棒性有限。",
                    "在 ROS 视觉应用中，它可以用于识别人脸、杯子或特定训练目标，再结合控制逻辑完成跟随或交互。",
                ],
            ),
        ],
        "review": [
            "能说明 ROS 图像到 OpenCV 图像的转换流程。",
            "能解释 HSV 为什么适合颜色检测。",
            "能描述目标跟随的闭环控制链路。",
            "能理解 Haar 级联分类器的基本思想。",
        ],
    },
    {
        "title": "八、机器人三维视觉仿真实例",
        "lead": "三维视觉关注物体在空间中的位置和形状，常用点云表示。PCL 是处理点云的核心库。",
        "sections": [
            (
                "PCL 库的作用和地位",
                [
                    "PCL 全称 Point Cloud Library，是点云处理领域常用的开源库。",
                    "点云由大量三维点组成，每个点通常包含 x、y、z 坐标，也可能包含颜色、法向量、强度等信息。",
                    "PCL 提供滤波、分割、特征提取、配准、表面重建、目标识别等算法。",
                    "在 ROS 中，点云常用 `sensor_msgs/PointCloud2` 消息表示，PCL 可以与 ROS 点云消息互相转换。",
                    "PCL 对三维感知很重要，例如机器人识别桌面、检测杯子、估计物体位姿、构建三维地图。",
                ],
            ),
            (
                "本章涉及的 PCL 主要算法流程",
                [
                    "数据获取：从深度相机、RGB-D 相机或三维激光雷达获取点云。",
                    "格式转换：把 ROS 的 PointCloud2 转为 PCL 点云类型，便于调用 PCL 算法。",
                    "滤波预处理：使用直通滤波限制空间范围，使用体素滤波降低点数，使用统计滤波去除离群噪声。",
                    "平面分割：常用 RANSAC 找出桌面、墙面、地面等主要平面，并把平面点和物体点分离。",
                    "聚类分割：对剩余点云做欧式聚类，把不同物体分成不同点云簇。",
                    "特征提取：计算法向量、边界、FPFH 等特征，用于识别或配准。",
                    "配准：把不同视角或不同时间的点云对齐，常见算法有 ICP。",
                    "结果发布与可视化：把处理后的点云、目标位置或识别结果发布回 ROS，在 Rviz 中查看。",
                ],
            ),
        ],
        "review": [
            "能解释点云是什么，以及 PCL 解决什么问题。",
            "能按顺序说出点云处理基本流程：获取、转换、滤波、分割、识别/配准、发布。",
        ],
    },
    {
        "title": "九、基于 ROS 的服务机器人应用",
        "lead": "服务机器人应用通常不是单一算法，而是感知、决策、导航、机械臂控制和人机交互组合而成。状态机是组织复杂任务流程的常用思想。",
        "sections": [
            (
                "状态机编程思想",
                [
                    "状态机把复杂任务拆成有限个状态，每个状态完成一件明确的事，并根据条件跳转到下一个状态。",
                    "例如饮料抓取任务可以拆成：等待命令、导航到饮料区、识别饮料、靠近目标、机械臂抓取、导航到用户、放置饮料、任务结束。",
                    "状态机的好处是逻辑清晰、便于调试、便于处理失败和重试。",
                    "每个状态通常有入口动作、执行动作、退出条件和异常处理。比如抓取失败可以重新识别或返回等待人工处理。",
                    "ROS 中可以用程序手写状态机，也可以用 SMACH 等状态机框架组织任务。",
                ],
            ),
            (
                "饮料抓取和放置的算法流程",
                [
                    "接收任务：通过语音、按钮、服务调用或上层指令获得饮料需求。",
                    "导航到目标区域：调用 move_base 或导航动作，到达饮料所在位置。",
                    "感知识别：使用相机、深度相机或点云识别饮料位置、类别和抓取姿态。",
                    "目标定位：把视觉坐标转换到机器人或机械臂基坐标系，需要 TF 坐标变换。",
                    "路径规划：机械臂规划从当前姿态到预抓取姿态、抓取姿态和抬起姿态的轨迹。",
                    "执行抓取：控制夹爪闭合，检测是否抓稳。",
                    "移动运输：底盘导航到放置区域或用户位置。",
                    "放置饮料：机械臂移动到放置位，打开夹爪，撤回机械臂。",
                    "任务反馈：向用户或上层系统报告成功、失败或异常原因。",
                ],
            ),
            (
                "核心技术",
                [
                    "导航定位：让机器人知道自己在哪里，并能到达目标区域。",
                    "视觉识别：识别饮料、杯子、桌面和障碍物。",
                    "三维定位与 TF：把相机看到的目标位置转换到机械臂可执行的坐标系。",
                    "机械臂运动规划：生成安全、可达、无碰撞的机械臂轨迹。",
                    "抓取控制：控制夹爪力度和开合，避免抓不住或损坏物体。",
                    "状态机调度：把多个模块串成稳定流程，并处理失败重试。",
                    "人机交互：接收用户需求、显示状态、语音提示或任务确认。",
                ],
            ),
        ],
        "review": [
            "能用状态机描述服务机器人任务。",
            "能说明饮料抓取放置流程涉及导航、视觉、TF、机械臂和控制。",
        ],
    },
]


tables = [
    {
        "title": "ROS 四种通信方式对比",
        "headers": ["方式", "同步/异步", "数据特点", "典型场景", "例子"],
        "rows": [
            ["Topic", "异步", "连续数据流，多对多", "传感器、速度、状态", "/scan、/cmd_vel"],
            ["Service", "同步", "一次请求一次响应", "查询、触发短任务", "保存地图、加法服务"],
            ["Action", "异步", "目标、反馈、结果，可取消", "耗时任务", "导航到目标点"],
            ["Parameter", "按需读写", "全局配置键值", "阈值、频率、文件路径", "/robot_name"],
        ],
    },
    {
        "title": "Gazebo 与 Rviz 对比",
        "headers": ["工具", "核心作用", "是否物理仿真", "常见用途"],
        "rows": [
            ["Gazebo", "构建虚拟世界并模拟机器人运动", "是", "仿真环境、传感器、碰撞、动力学"],
            ["Rviz", "显示 ROS 数据和机器人状态", "否", "查看 TF、地图、路径、点云、模型"],
        ],
    },
    {
        "title": "Navigation 关键功能包",
        "headers": ["功能包", "作用", "输入/依赖", "输出/结果"],
        "rows": [
            ["gmapping", "SLAM 建图", "激光、里程计、TF", "二维栅格地图"],
            ["map_server", "地图加载与保存", "地图 yaml/pgm", "map 话题"],
            ["amcl", "已有地图中定位", "地图、激光、里程计、TF", "机器人位姿估计"],
            ["move_base", "导航协调与速度输出", "目标点、地图、传感器、TF", "/cmd_vel"],
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("嵌入式系统原理复习纲要完全版知识体系")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def style_document(doc):
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 20, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "365F91"),
    ]:
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_toc_placeholder(doc):
    p = doc.add_paragraph()
    run = p.add_run("目录提示：在 Word 中打开后，可在“引用 - 目录”中插入或更新自动目录。")
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_bullets(doc, bullets):
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_docx():
    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026 春嵌入式系统原理复习纲要\n完全版知识体系")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("按原复习纲要扩展整理，面向零基础复习使用").italic = True
    doc.add_paragraph()
    doc.add_heading("使用说明", level=1)
    add_bullets(
        doc,
        [
            "本资料覆盖原纲要中列出的全部知识点，并补充必要背景、概念解释、命令、流程和对比。",
            "零基础阅读顺序建议：先理解 Linux 与嵌入式基础，再学习 ROS 通信，最后看导航、视觉和服务机器人应用。",
            "考试复习时优先掌握每章末尾的“本章应会”。如果能用自己的话解释这些条目，说明已经建立了主干知识体系。",
        ],
    )
    add_toc_placeholder(doc)

    doc.add_heading("总览：本课程知识主线", level=1)
    add_bullets(
        doc,
        [
            "第一层是嵌入式与 Linux：理解资源受限设备如何运行，掌握命令行与软件安装。",
            "第二层是 ROS 基础：理解工作空间、功能包、节点和编译运行流程。",
            "第三层是 ROS 通信：掌握主题、服务、动作和参数服务器，这是后续所有机器人应用的基础。",
            "第四层是工具与建模：使用 TF、launch、Gazebo、Rviz、URDF 搭建和观察机器人系统。",
            "第五层是应用算法：建图导航、二维视觉、三维点云和服务机器人任务调度。",
        ],
    )

    for chapter in chapters:
        doc.add_heading(chapter["title"], level=1)
        doc.add_paragraph(chapter["lead"])
        for title, bullets in chapter["sections"]:
            doc.add_heading(title, level=2)
            add_bullets(doc, bullets)
        doc.add_heading("本章应会", level=2)
        add_bullets(doc, chapter["review"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("核心对比表", level=1)
    for t in tables:
        doc.add_heading(t["title"], level=2)
        table = doc.add_table(rows=1, cols=len(t["headers"]))
        table.style = "Table Grid"
        hdr = table.rows[0]
        set_repeat_table_header(hdr)
        for idx, head in enumerate(t["headers"]):
            hdr.cells[idx].text = head
            set_cell_shading(hdr.cells[idx], "D9EAF7")
            for p in hdr.cells[idx].paragraphs:
                for run in p.runs:
                    run.bold = True
        for row in t["rows"]:
            cells = table.add_row().cells
            for idx, val in enumerate(row):
                cells[idx].text = val
        doc.add_paragraph()

    doc.add_heading("最后复习路线", level=1)
    add_bullets(
        doc,
        [
            "先背主线：嵌入式特点 - Linux - ROS 工作空间 - ROS 通信 - TF/launch - 仿真建模 - 导航 - 视觉 - 服务机器人。",
            "再练对比：Topic/Service/Action/Parameter，Gazebo/Rviz，全局规划/局部规划，RGB/HSV。",
            "最后练流程：主题发布订阅流程，自定义服务流程，launch 启动流程，SLAM 与 Navigation 流程，视觉目标跟随流程，饮料抓取放置状态机流程。",
            "遇到名词题时先答定义，再答作用，再给例子。遇到流程题时按输入、处理、输出三个层次组织答案。",
        ],
    )
    doc.save(DOCX_PATH)


def build_markdown():
    lines = []
    lines.append("# 2026 春嵌入式系统原理复习纲要 - 完全版知识体系")
    lines.append("")
    lines.append("> 按原复习纲要扩展整理，面向零基础复习使用。")
    lines.append("")
    lines.append("## 使用说明")
    lines.extend(
        [
            "- 本资料覆盖原纲要中列出的全部知识点，并补充必要背景、概念解释、命令、流程和对比。",
            "- 零基础阅读顺序建议：先理解 Linux 与嵌入式基础，再学习 ROS 通信，最后看导航、视觉和服务机器人应用。",
            "- 考试复习时优先掌握每章末尾的“本章应会”。",
            "",
            "## 总览：本课程知识主线",
            "- 第一层是嵌入式与 Linux：理解资源受限设备如何运行，掌握命令行与软件安装。",
            "- 第二层是 ROS 基础：理解工作空间、功能包、节点和编译运行流程。",
            "- 第三层是 ROS 通信：掌握主题、服务、动作和参数服务器。",
            "- 第四层是工具与建模：使用 TF、launch、Gazebo、Rviz、URDF 搭建和观察机器人系统。",
            "- 第五层是应用算法：建图导航、二维视觉、三维点云和服务机器人任务调度。",
            "",
        ]
    )
    for chapter in chapters:
        lines.append(f"## {chapter['title']}")
        lines.append("")
        lines.append(chapter["lead"])
        lines.append("")
        for title, bullets in chapter["sections"]:
            lines.append(f"### {title}")
            lines.append("")
            for item in bullets:
                lines.append(f"- {item}")
            lines.append("")
        lines.append("### 本章应会")
        lines.append("")
        for item in chapter["review"]:
            lines.append(f"- {item}")
        lines.append("")
    lines.append("## 核心对比表")
    lines.append("")
    for t in tables:
        lines.append(f"### {t['title']}")
        lines.append("")
        lines.append("| " + " | ".join(t["headers"]) + " |")
        lines.append("| " + " | ".join(["---"] * len(t["headers"])) + " |")
        for row in t["rows"]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    lines.append("## 最后复习路线")
    lines.append("")
    lines.extend(
        [
            "- 先背主线：嵌入式特点 - Linux - ROS 工作空间 - ROS 通信 - TF/launch - 仿真建模 - 导航 - 视觉 - 服务机器人。",
            "- 再练对比：Topic/Service/Action/Parameter，Gazebo/Rviz，全局规划/局部规划，RGB/HSV。",
            "- 最后练流程：主题发布订阅流程，自定义服务流程，launch 启动流程，SLAM 与 Navigation 流程，视觉目标跟随流程，饮料抓取放置状态机流程。",
            "- 遇到名词题时先答定义，再答作用，再给例子。遇到流程题时按输入、处理、输出三个层次组织答案。",
        ]
    )
    MD_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_markdown()
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
