#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/cuing/ros/ros-course-lab")
OUT = ROOT / "docs/project_documentation/mountain_car_sim_implementation_guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_para(doc, text="", style=None, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        set_east_asia_font(run)
        rest = text[len(bold_prefix):]
        if rest:
            run = paragraph.add_run(rest)
            set_east_asia_font(run)
    else:
        run = paragraph.add_run(text)
        set_east_asia_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        p = add_para(doc, item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_numbers(doc, items):
    for item in items:
        p = add_para(doc, item, style="List Number")
        p.paragraph_format.space_after = Pt(4)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="E5E7EB", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)
    return table


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="C9D8EA", size="6")
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, "F3F7FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    set_east_asia_font(run)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run(body)
    set_east_asia_font(run)
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        set_east_asia_font(run)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            set_east_asia_font(run)
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    run = p.add_run("mountain_car_sim 技术实现文档")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = DARK_BLUE
    set_east_asia_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROS Noetic + Gazebo11 山地场景与自动导航小车")
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED
    set_east_asia_font(run)

    doc.add_paragraph()
    add_table(
        doc,
        ["项目项", "内容"],
        [
            ["实验名称", "新场景搭建1"],
            ["ROS 包名", "mountain_car_sim"],
            ["目标环境", "Ubuntu 20.04 / ROS Noetic / Gazebo Classic 11"],
            ["核心功能", "山地场景导入、四轮小车建模、WASD 控制、自动寻迹避障到红旗"],
        ],
        [2300, 7060],
    )
    add_note(
        doc,
        "文档用途",
        "本文按“从 0 开始实现”的顺序组织，可直接作为实验报告的技术实现部分，也可用于现场答辩时解释项目架构和运行流程。",
    )
    doc.add_page_break()


def add_manual_toc(doc):
    doc.add_heading("目录", level=1)
    items = [
        "1. 项目目标与最终效果",
        "2. 开发环境准备",
        "3. 从 0 创建 ROS 包结构",
        "4. 山地场景搭建",
        "5. 四轮小车模型设计",
        "6. Gazebo 导入与启动文件",
        "7. 控制脚本：WASD、自动巡航与自动导航",
        "8. 编译、运行与演示流程",
        "9. 常见问题排查",
        "10. 答辩截图建议与总结",
    ]
    add_numbers(doc, items)


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_manual_toc(doc)

    doc.add_heading("1. 项目目标与最终效果", level=1)
    add_para(
        doc,
        "本项目面向 ROS Noetic 与 Gazebo Classic 11 课程实验，目标是在一个完整 catkin 工作空间中创建可运行的山地小车仿真包。项目最终可以通过一条 roslaunch 命令启动山地场景、导入小车，并让小车沿山路绕过岩石障碍物后到达红旗检查点。",
    )
    add_bullets(
        doc,
        [
            "场景层：包含程序化山地 mesh、蛇形土路、岩石障碍物、红旗检查点、太阳光和默认相机视角。",
            "机器人层：使用 xacro 编写四轮小车，包含底盘、车顶识别块、四个车轮、visual/collision/inertial 和 Gazebo 控制插件。",
            "控制层：提供 /cmd_vel 速度控制接口，支持 WASD 手动控制、简单自动巡航和自动寻迹避障。",
            "展示层：默认 demo.launch 会启动自动寻迹避障，适合直接录屏或截图用于实验报告。",
        ],
    )

    doc.add_heading("2. 开发环境准备", level=1)
    add_para(doc, "推荐环境为 Ubuntu 20.04、ROS Noetic、Gazebo Classic 11。进入工作空间后，先加载 ROS 环境变量，再确认 catkin 工作空间结构存在。")
    add_code(
        doc,
        "source /opt/ros/noetic/setup.bash\n"
        "cd ~/ros-course-lab\n"
        "ls src\n",
    )
    add_note(
        doc,
        "编译提示",
        "如果同一个工作空间里还有其他课程包，并且其他包存在缺失依赖，可以使用 CATKIN_WHITELIST_PACKAGES 只编译 mountain_car_sim，避免无关包阻塞本实验。",
    )
    add_code(
        doc,
        "catkin_make -DCATKIN_WHITELIST_PACKAGES=mountain_car_sim\n"
        "source devel/setup.bash\n",
    )

    doc.add_heading("3. 从 0 创建 ROS 包结构", level=1)
    add_para(doc, "从空工作空间开始时，先在 src 目录下创建 ROS 包，再补齐 launch、worlds、urdf、models、scripts 等目录。")
    add_code(
        doc,
        "cd ~/ros-course-lab/src\n"
        "catkin_create_pkg mountain_car_sim rospy std_msgs geometry_msgs gazebo_ros gazebo_msgs xacro robot_state_publisher\n"
        "cd mountain_car_sim\n"
        "mkdir -p launch worlds urdf models scripts\n",
    )
    add_table(
        doc,
        ["路径", "作用"],
        [
            ["launch/", "保存 Gazebo 场景、小车生成和一键 demo 的 roslaunch 文件。"],
            ["worlds/mountain_scene.world", "定义山地世界、光照、物理参数、模型 include 和默认相机视角。"],
            ["urdf/mountain_car.xacro", "定义四轮小车结构、惯量、碰撞体和 Gazebo 控制插件。"],
            ["models/mountain_terrain/", "保存山地地形模型、程序化 DAE mesh 和模型配置。"],
            ["models/rock_obstacle/", "保存岩石障碍物 SDF 模型。"],
            ["models/checkpoint_flag/", "保存红旗检查点 SDF 模型。"],
            ["scripts/", "保存生成模型、WASD 控制和自动寻迹避障节点。"],
        ],
        [3300, 6060],
    )

    doc.add_heading("4. 山地场景搭建", level=1)
    add_para(
        doc,
        "早期版本如果只用几个 box 表示山坡，会显得像普通几何块。当前版本将山地主体改为 DAE 网格模型，地形文件位于 models/mountain_terrain/meshes/，并在 SDF 中通过 model:// URI 引用。",
    )
    add_bullets(
        doc,
        [
            "mountain_terrain.dae：山地主体 mesh，包含主坡度、两侧山脊和小幅起伏。",
            "mountain_trail.dae：贴合地形高度的蛇形土路，仅作为视觉层，帮助展示路线。",
            "terrain_base：山地外圈底座，使场地从 Gazebo 网格背景中更容易辨认。",
            "left_boundary / right_boundary：低矮护栏，用于强调场地边界并降低演示跑出场景的概率。",
        ],
    )
    add_para(doc, "worlds/mountain_scene.world 中通过 include 加载模型，并设置物理参数、环境光、太阳光和默认相机。")
    add_code(
        doc,
        "<include>\n"
        "  <uri>model://mountain_terrain</uri>\n"
        "  <pose>0 0 0 0 0 0</pose>\n"
        "</include>\n"
        "<include>\n"
        "  <name>rock_obstacle_left</name>\n"
        "  <uri>model://rock_obstacle</uri>\n"
        "  <pose>-1.2 1.15 0.18 0 0 0.35</pose>\n"
        "</include>\n"
        "<include>\n"
        "  <uri>model://checkpoint_flag</uri>\n"
        "  <pose>4.7 0.0 1.05 0 0 0</pose>\n"
        "</include>\n",
    )

    doc.add_heading("5. 四轮小车模型设计", level=1)
    add_para(
        doc,
        "小车模型使用 xacro 编写，优势是尺寸、质量、车轮半径等参数可以集中管理。模型包含 base_footprint、base_link、top_cabin 和四个车轮 link。",
    )
    add_table(
        doc,
        ["组成", "实现要点"],
        [
            ["base_footprint", "作为地面投影坐标系，避免根 link 直接带惯量导致 KDL 警告。"],
            ["base_link", "橙色车体，包含 box visual、collision 和质量/惯量。"],
            ["top_cabin", "蓝色车顶块，用来区分车头方向，便于截图观察。"],
            ["四个 wheel link", "圆柱体车轮，visual/collision 均旋转 90 度，使车轮轴线沿 y 轴。"],
            ["Gazebo 接触参数", "对车轮设置 mu、mu2、kp、kd，减少抖动和穿模风险。"],
            ["planar_move 插件", "订阅 /cmd_vel 并控制车体移动，适合课程演示环境稳定运行。"],
        ],
        [2400, 6960],
    )
    add_code(
        doc,
        "<gazebo>\n"
        "  <plugin name=\"mountain_car_planar_move\" filename=\"libgazebo_ros_planar_move.so\">\n"
        "    <commandTopic>cmd_vel</commandTopic>\n"
        "    <odometryTopic>odom</odometryTopic>\n"
        "    <robotBaseFrame>base_footprint</robotBaseFrame>\n"
        "    <publishTF>true</publishTF>\n"
        "  </plugin>\n"
        "</gazebo>\n",
    )

    doc.add_heading("6. Gazebo 导入与启动文件", level=1)
    add_para(doc, "项目使用三个 launch 文件分层组织，便于单独调试场景、小车和完整 demo。")
    add_table(
        doc,
        ["文件", "关键职责"],
        [
            ["mountain_world.launch", "设置 GAZEBO_MODEL_PATH，调用 gazebo_ros/empty_world.launch 加载 mountain_scene.world。"],
            ["spawn_mountain_car.launch", "用 xacro 生成 robot_description，再调用 reset_spawn_model.py 生成小车。"],
            ["demo.launch", "一键启动场景、小车和控制节点，默认 controller_mode=flag_nav。"],
        ],
        [3100, 6260],
    )
    add_para(doc, "重复运行 launch 时，Gazebo 里如果已有同名模型，会出现 entity already exists。项目通过 reset_spawn_model.py 先调用 /gazebo/delete_model 删除旧模型，再调用 /gazebo/spawn_urdf_model 生成新模型。")

    doc.add_heading("7. 控制脚本：WASD、自动巡航与自动导航", level=1)
    doc.add_heading("7.1 /cmd_vel 控制接口", level=2)
    add_para(doc, "小车控制统一使用 geometry_msgs/Twist 发布到 /cmd_vel。线速度 linear.x 控制前进/后退，角速度 angular.z 控制左转/右转。")
    doc.add_heading("7.2 WASD 手动控制", level=2)
    add_para(doc, "simple_car_controller.py 的 keyboard 模式从终端读取按键，因此必须在单独终端中 rosrun，不能依赖 roslaunch 窗口读取键盘输入。")
    add_code(
        doc,
        "rosrun mountain_car_sim simple_car_controller.py _mode:=keyboard\n"
        "# w 前进，s 后退，a 左转，d 右转，空格停止，q 退出\n",
    )
    doc.add_heading("7.3 自动寻迹避障", level=2)
    add_para(
        doc,
        "autonomous_flag_nav.py 订阅 /gazebo/model_states 获取 mountain_car、rock_obstacle 和 checkpoint_flag 的位置。控制器按蛇形山路航点前进，同时对附近岩石生成斥力向量，最后发布速度到 /cmd_vel。",
    )
    add_table(
        doc,
        ["步骤", "说明"],
        [
            ["读取状态", "从 /gazebo/model_states 获取小车姿态、岩石位置和红旗位置。"],
            ["选择航点", "根据当前位置在预设山路航点列表中切换目标。"],
            ["计算偏航", "由目标方向与小车 yaw 的差值得到 heading_error。"],
            ["障碍物斥力", "当岩石距离小于 avoid_radius 时，计算远离岩石的 repulsive vector。"],
            ["发布速度", "线速度随目标距离和偏航角调整，角速度按 heading_error 限幅。"],
            ["终点停止", "进入 goal_tolerance 后持续发布零速度。"],
        ],
        [2100, 7260],
    )
    add_code(
        doc,
        "目标方向 = 航点方向 + 障碍物斥力\n"
        "angular.z = clamp(1.8 * heading_error, -max_turn, max_turn)\n"
        "linear.x  = clamp(0.75 * target_dist, min_speed, max_speed) * cos_factor\n",
    )

    doc.add_heading("8. 编译、运行与演示流程", level=1)
    add_para(doc, "最小运行流程如下。若工作空间只有本项目，可以直接 catkin_make；若有其他缺失依赖包，建议使用白名单编译。")
    add_code(
        doc,
        "cd ~/ros-course-lab\n"
        "source /opt/ros/noetic/setup.bash\n"
        "catkin_make -DCATKIN_WHITELIST_PACKAGES=mountain_car_sim\n"
        "source devel/setup.bash\n"
        "roslaunch mountain_car_sim demo.launch\n",
    )
    add_para(doc, "常用变体命令：")
    add_code(
        doc,
        "# 只启动山地世界\n"
        "roslaunch mountain_car_sim mountain_world.launch\n\n"
        "# 启动场景和小车，但不自动控制\n"
        "roslaunch mountain_car_sim demo.launch start_controller:=false\n\n"
        "# 单独启动 WASD 手动控制\n"
        "rosrun mountain_car_sim simple_car_controller.py _mode:=keyboard\n\n"
        "# 使用旧版自动巡航\n"
        "roslaunch mountain_car_sim demo.launch controller_mode:=auto\n",
    )

    doc.add_heading("9. 常见问题排查", level=1)
    add_table(
        doc,
        ["现象", "原因", "处理方式"],
        [
            ["Gazebo 找不到模型", "GAZEBO_MODEL_PATH 未包含本包 models 目录。", "使用 mountain_world.launch 启动，或检查 source devel/setup.bash 是否执行。"],
            ["SpawnModel: entity already exists", "Gazebo 中已有 mountain_car。", "使用当前 spawn_mountain_car.launch，它会先 delete 再 spawn；必要时重启 Gazebo。"],
            ["按 WASD 没反应", "键盘焦点不在运行 rosrun 的终端，或控制节点不是 keyboard 模式。", "新开终端运行 rosrun mountain_car_sim simple_car_controller.py _mode:=keyboard。"],
            ["小车不动", "/cmd_vel 没有发布或 Gazebo 插件未加载。", "用 rostopic echo /cmd_vel 检查速度，用 rostopic list 检查 /odom。"],
            ["catkin_make 被其他包卡住", "同工作空间其他课程包有缺失依赖。", "使用 catkin_make -DCATKIN_WHITELIST_PACKAGES=mountain_car_sim。"],
            ["小车抖动或穿模", "坡面接触和物理步长不稳定。", "保留 world 中较小 max_step_size 和车轮接触参数，避免把速度调得过大。"],
        ],
        [2100, 3350, 3910],
    )

    doc.add_heading("10. 答辩截图建议与总结", level=1)
    add_para(doc, "答辩时建议按“场景、模型、控制、结果”四类截图组织，能覆盖课程要求，也能让老师快速看到项目完成度。")
    add_bullets(
        doc,
        [
            "山地场景整体图：从 Gazebo 默认相机视角截取，包含山体、蛇形土路、岩石和红旗。",
            "小车模型导入图：选中 mountain_car，显示橙色车身、蓝色车顶和四个黑色车轮。",
            "小车运动图：截取小车沿山路靠近或绕过岩石时的画面。",
            "到达红旗图：截取小车停在 checkpoint_flag 附近的最终画面。",
            "终端运行图：展示 catkin_make、source、roslaunch mountain_car_sim demo.launch 等命令。",
        ],
    )
    add_note(
        doc,
        "项目总结",
        "该项目完成了从 ROS 包创建、Gazebo 场景建模、URDF/xacro 机器人建模、模型导入、/cmd_vel 控制到自动导航演示的闭环。后续如果继续提升，可以加入真实传感器、激光雷达避障、导航栈或更精细的 DEM 地形。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
